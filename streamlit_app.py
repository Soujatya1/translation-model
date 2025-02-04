import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from deep_translator import GoogleTranslator
import concurrent.futures

def translate_run(run_text, translator):
    """
    Helper function to translate individual run text.
    :param run_text: The text of a run to translate.
    :param translator: The Google Translator instance.
    :return: Translated text.
    """
    try:
        translated_text = translator.translate(run_text.strip()) or run_text
        return translated_text
    except Exception as e:
        print(f"Error translating run: {e}")
        return run_text  # Return original text in case of an error

def translate_doc(doc, destination='hi'):
    """
    Translate a Word document while preserving formatting efficiently using concurrency.
    :param doc: Word doc object (from Document class)
    :param destination: Target language (default is Hindi 'hi')
    """
    translator = GoogleTranslator(source='auto', target=destination)

    # Create a ThreadPoolExecutor for parallel translation
    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = []

        # Translate paragraphs and runs concurrently
        for p in doc.paragraphs:
            if p.text.strip():
                for run in p.runs:
                    if run.text.strip():
                        futures.append(executor.submit(translate_run, run.text.strip(), translator))

        # Translate table cells and runs concurrently
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text.strip():
                                    futures.append(executor.submit(translate_run, run.text.strip(), translator))

        # Wait for all translations to finish and update runs with translated text
        for future in concurrent.futures.as_completed(futures):
            translated_text = future.result()
            # Find the corresponding run and update the text (this step assumes text is updated in place)
            for p in doc.paragraphs:
                for run in p.runs:
                    if run.text.strip() == translated_text:
                        run.text = translated_text

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text.strip() == translated_text:
                                    run.text = translated_text

    return doc
def main():
    st.title("Word Document Translator")
    
    uploaded_file = st.file_uploader("Upload a Word Document", type=["docx"])
    if uploaded_file:
        doc = Document(uploaded_file)
        
        language_options = {
            "Bengali": "bn", "Hindi": "hi", "Odia": "or", "Punjabi": "pa", 
            "Tamil": "ta", "Telegu": "te", "Gujarati": "gu", "Malayalam": "ml"
        }
        target_language = st.selectbox("Select Target Language", options=list(language_options.keys()))
        language_code = language_options[target_language]
        
        if st.button("Translate Document"):
            with st.spinner('Translating...'):
                translated_doc = translate_doc(doc, language_code)
                
                with open("translated_document.docx", "wb") as f:
                    translated_doc.save(f)
                
                with open("translated_document.docx", "rb") as f:
                    st.download_button(
                        label="Download Translated Document",
                        data=f,
                        file_name="translated_document.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                st.success("Translation complete!")

if __name__ == '__main__':
    main()
